from datetime import datetime
from bs4 import BeautifulSoup as bs
import unirest
import scraperwiki
import re
from docx import Document
from docx.shared import Inches

user_agent = {'User-Agent': 'Mozilla/5.0 (compatible; bingbot/2.0; +http://www.bing.com/bingbot.htm)'}
document = Document()
pages = unirest.get('http://www.foodnetwork.com/recipes/food-network-kitchens/roasted-cauliflower-recipe.html', headers = user_agent) # make a request to the url

soup = bs(pages.raw_body)
title = soup.find('h1', attrs ={'itemprop':'name'}).text
document.add_heading(title, 0)
time  = soup.find( text=re.compile('Total Time:'))
t = time.find_next('dd').text
document.add_paragraph('Total Time: ' +t)
yields = soup.find( text=re.compile('Yield:'))
y = yields.find_next('dd').text
document.add_paragraph('Yield: ' +y)


document.save('demo.docx')
